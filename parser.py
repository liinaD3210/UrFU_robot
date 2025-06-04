# parser.py

import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import os

BASE_URL = "https://priem-rtf-old.urfu.ru"
START_URL = f"{BASE_URL}/ru/baccalaureate/"
OUTPUT_DOCX_FILE = "data/irit_rtf_baccalaureate_info.docx"

HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
}

def get_soup(url):
    """Получает HTML-контент по URL и возвращает объект BeautifulSoup."""
    print(f"DEBUG: Запрос к URL: {url}")
    try:
        response = requests.get(url, headers=HEADERS, timeout=15)
        print(f"DEBUG: Статус код ответа от {url}: {response.status_code}")
        response.raise_for_status()
        
        # Раскомментируйте следующую строку, чтобы сохранить HTML в файл для анализа
        with open("debug_page_content.html", "w", encoding=response.encoding or 'utf-8') as f:
             f.write(response.text)
        print(f"DEBUG: HTML контент со страницы {url} сохранен в debug_page_content.html")

        soup = BeautifulSoup(response.content, 'html.parser')
        # print(f"DEBUG: BeautifulSoup объект создан. Первые 500 символов prettify для {url}:")
        # print(soup.prettify()[:500])
        return soup
    except requests.exceptions.RequestException as e:
        print(f"ОШИБКА при запросе к {url}: {e}")
        return None

def extract_program_links(main_page_soup):
    """Извлекает названия направлений и ссылки на их страницы."""
    programs = []
    
    # Ищем контейнер с программами бакалавриата
    container_bacalavr = main_page_soup.select_one('section.containerBacalavr')
    
    if not container_bacalavr:
        print("ПРЕДУПРЕЖДЕНИЕ: Контейнер 'section.containerBacalavr' не найден на странице.")
        # Дополнительная отладка: попробуем найти другие характерные элементы, если основной не найден
        tab_och = main_page_soup.select_one('label[for="tab1"]')
        if tab_och:
            print(f"DEBUG: Найдена вкладка '{tab_och.get_text(strip=True)}', значит, часть структуры есть.")
        else:
            print("DEBUG: Даже вкладка 'Очная форма обучения' не найдена. Проверьте HTML в debug файле.")
        return []

    # Внутри контейнера ищем все блоки с классом 'oneBlock'
    program_blocks = container_bacalavr.select('div.oneBlock')
    
    print(f"DEBUG: Найдено блоков программ ('div.oneBlock'): {len(program_blocks)}")

    if not program_blocks:
        print("ПРЕДУПРЕЖДЕНИЕ: Блоки 'div.oneBlock' не найдены внутри 'section.containerBacalavr'.")
        return []

    for i, block in enumerate(program_blocks):
        link_tag = block.find('a') # Ссылка должна быть прямо внутри oneBlock
        
        if link_tag and link_tag.has_attr('href'):
            relative_url = link_tag['href']
            # Убедимся, что URL абсолютный или корректно соединяется с BASE_URL
            if relative_url.startswith('http'):
                absolute_url = relative_url
            else:
                absolute_url = urljoin(BASE_URL, relative_url) # BASE_URL = "https://priem-rtf-old.urfu.ru"

            # Ищем название внутри div.blockOfInformation > h2
            info_block = link_tag.find('div', class_='blockOfInformation')
            name = "Название не найдено"
            if info_block:
                name_tag = info_block.find('h2')
                if name_tag:
                    # Удаляем вложенный span с кодом направления, если он есть, и лишние пробелы/переносы
                    for span in name_tag.find_all("span"):
                        span.decompose() # Удаляем тег span и его содержимое
                    name = name_tag.get_text(separator=' ', strip=True).replace('\n', ' ').replace('  ', ' ')

            programs.append({'name': name, 'url': absolute_url})
            print(f"Найдено направление: {name} - {absolute_url}")
        else:
            print(f"ПРЕДУПРЕЖДЕНИЕ: В блоке {i+1} ('div.oneBlock') не удалось найти тег 'a' с атрибутом 'href'.")
            # print(f"DEBUG: HTML проблемного блока:\n{block.prettify()}") # Раскомментировать для детального просмотра блока
    
    if not programs and program_blocks:
         print("ПРЕДУПРЕЖДЕНИЕ: Блоки 'div.oneBlock' были найдены, но из них не удалось извлечь информацию о программах (название/ссылку). Проверьте внутреннюю структуру блоков.")

    return programs

def extract_program_description(program_url):
    """Извлекает описание программы с ее индивидуальной страницы."""
    print(f"  Парсинг описания для: {program_url}")
    soup = get_soup(program_url) # get_soup уже включает отладку и сохранение HTML
    if not soup:
        return "Не удалось загрузить описание (soup is None)."

    description_parts = []
    
    # Новый селектор для основного контента (описания)
    # Ищем <div class="ce-bodytext">
    content_area = soup.select_one('div.ce-bodytext')

    if content_area:
        # print(f"DEBUG: Найдена контентная область 'div.ce-bodytext' для {program_url}.")
        
        # Извлекаем текст из всех тегов <p> внутри этой области
        paragraphs = content_area.find_all('p')
        for p_tag in paragraphs:
            # Можно добавить дополнительную обработку текста, если нужно
            # Например, убрать ссылки или специфические фразы
            
            # Копируем тег, чтобы не изменять оригинальный soup для извлечения текста
            temp_p_tag = BeautifulSoup(str(p_tag), 'html.parser').p 
            
            # Удаляем все ссылки <a> из параграфа перед извлечением текста,
            # чтобы избежать текста ссылок типа "проектный практикум" в описании,
            # если мы хотим только чистый текст описания.
            if temp_p_tag:
                for a_tag in temp_p_tag.find_all('a'):
                    a_tag.decompose() # Удаляет тег <a> и его содержимое

                text = temp_p_tag.get_text(separator=' ', strip=True)
            else:
                text = p_tag.get_text(separator=' ', strip=True)


            if text and len(text) > 15: # Простая фильтрация коротких/пустых параграфов
                description_parts.append(text)
        
        if not description_parts:
            print(f"    ПРЕДУПРЕЖДЕНИЕ: Не найдено подходящих тегов <p> с текстом внутри 'div.ce-bodytext' на {program_url}.")
            # Если <p> не дали результата, можно попробовать взять весь текст из content_area,
            # но это может захватить лишнее, если структура сложнее.
            # all_text = content_area.get_text(separator='\n', strip=True)
            # if all_text and len(all_text) > 20:
            #     description_parts.append(all_text)
            # else:
            #     return "Описание не найдено (альтернативный метод тоже не сработал)."
            return "Описание не найдено (теги <p> внутри 'div.ce-bodytext' не содержат подходящей информации или пусты)."
        
        return "\n\n".join(description_parts)
    else:
        print(f"    ПРЕДУПРЕЖДЕНИЕ: Не удалось найти контентную область ('div.ce-bodytext') на странице {program_url}")
        # Чтобы увидеть, что именно было загружено, можно раскомментировать сохранение HTML в get_soup
        # и посмотреть файл debug_page_content_НАЗВАНИЕПРОГРАММЫ.html
        # print(f"DEBUG: Содержимое soup для {program_url} (первые 1000 символов):\n{soup.prettify()[:1000]}")
        return "Структура страницы с описанием не распознана (отсутствует 'div.ce-bodytext')."

def save_to_docx(data, filename):
    """Сохраняет собранные данные в DOCX файл."""
    doc = Document()
    doc.add_heading('Направления бакалавриата ИРИТ-РТФ', level=0)

    for item in data:
        doc.add_heading(item['name'], level=1)
        
        description_paragraphs = item['description'].split('\n\n')
        for para_text in description_paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph() 

    os.makedirs(os.path.dirname(filename), exist_ok=True)
    doc.save(filename)
    print(f"Данные сохранены в файл: {filename}")


def main():
    print("Начало парсинга сайта ИРИТ-РТФ...")
    main_soup = get_soup(START_URL)
    
    if not main_soup:
        print("ОШИБКА: Не удалось загрузить главную страницу (main_soup is None). Выход.")
        return

    print("\nDEBUG: Запускаем извлечение ссылок на программы...")
    programs_info = extract_program_links(main_soup)
    
    if not programs_info:
        print("ОШИБКА: Направления не найдены или не удалось извлечь информацию о них. Проверьте DEBUG сообщения выше. Выход.")
        return

    print(f"\nDEBUG: Найдено {len(programs_info)} программ. Начинаем сбор описаний...")
    all_data = []
    for program in programs_info:
        description = extract_program_description(program['url'])
        all_data.append({
            'name': program['name'],
            'description': description
        })
        time.sleep(0.5) 

    if all_data:
        save_to_docx(all_data, OUTPUT_DOCX_FILE)
    else:
        print("ОШИБКА: Не удалось собрать данные для сохранения.")

if __name__ == "__main__":
    main()